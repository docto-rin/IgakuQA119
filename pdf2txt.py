from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import red, blue
from pdf2image import convert_from_path
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import json
import os
from docx import Document
from docx.shared import RGBColor
from dotenv import load_dotenv
from tqdm import tqdm
import time
import argparse

# .envファイルから環境変数を読み込む
load_dotenv()

# Azure Document Intelligence の設定
endpoint = os.getenv("AZURE_ENDPOINT")
key = os.getenv("AZURE_API_KEY")

if not endpoint or not key:
    raise ValueError("Azure credentials not found in .env file")

# BIZ UDPMincho フォントの登録
pdfmetrics.registerFont(TTFont("BIZUDPMincho", "fonts/BIZUDPMincho-Regular.ttf"))


def analyze_pdf(pdf_path):
    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )

    print("PDFを解析中...")
    try:
        with open(pdf_path, "rb") as pdf_file:
            # ファイルサイズを確認
            pdf_file.seek(0, 2)
            file_size = pdf_file.tell()
            pdf_file.seek(0)
            print(f"PDFファイルサイズ: {file_size / (1024 * 1024):.2f} MB")

            # APIリクエストを実行
            print("Azure Document Intelligenceにリクエストを送信...")
            poller = document_analysis_client.begin_analyze_document(
                "prebuilt-read", pdf_file
            )

            # ポーリング状態を監視
            print("処理中...")
            progress = tqdm(total=100, desc="OCR処理")
            while not poller.done():
                try:
                    # 進捗状況の更新（簡略化）
                    status = poller.status()
                    if status == "running":
                        progress.n = 50
                    elif status == "succeeded":
                        progress.n = 100
                    progress.refresh()
                except:
                    pass
                time.sleep(1)
            progress.close()
            print("処理完了")

            result = poller.result()

            print(f"合計 {len(result.pages)} ページを処理しました")
            for page_num, page in enumerate(result.pages, 1):
                word_count = len(page.words)
                print(f"  - ページ {page_num}: {word_count} 単語を検出")

            return result

    except Exception as e:
        print(f"エラーが発生しました: {str(e)}")
        if hasattr(e, "response"):
            print(f"レスポンスステータス: {e.response.status_code}")
            print(f"レスポンス内容: {e.response.text}")
        raise


def save_ocr_results(ocr_result, base_filename):
    # JSONファイルとして保存
    all_pages_data = []
    all_text_content = []

    print("OCR結果を保存中...")
    for page_num, page in tqdm(
        enumerate(ocr_result.pages, 1), total=len(ocr_result.pages), desc="ページ処理"
    ):
        page_data = {
            "page_number": page_num,
            "width": page.width,
            "height": page.height,
            "words": [
                {
                    "content": word.content,
                    "confidence": word.confidence,
                    "polygon": [(p.x, p.y) for p in word.polygon]
                    if word.polygon
                    else None,
                }
                for word in page.words
            ],
        }
        all_pages_data.append(page_data)

        # テキスト内容を収集（ページ単位）
        page_text = f"\n=== Page {page_num} ===\n"
        page_text += "".join(word.content for word in page.words)
        all_text_content.append(page_text)

    # JSON保存
    json_filename = f"{base_filename}_ocr_results.json"
    with open(json_filename, "w", encoding="utf-8") as f:
        json.dump(all_pages_data, f, ensure_ascii=False, indent=4)

    # テキストファイル保存
    text_filename = f"{base_filename}_ocr_text.txt"
    with open(text_filename, "w", encoding="utf-8") as f:
        f.write("\n".join(all_text_content))

    # Wordファイルの作成
    doc = Document()
    for page_num, page in enumerate(ocr_result.pages, 1):
        doc.add_heading(f"Page {page_num}", level=1)
        paragraph = doc.add_paragraph()

        # テキストファイルと同じ順番でワードを追加
        for word_obj in page.words:
            run = paragraph.add_run(word_obj.content)
            if float(word_obj.confidence) < 0.95:
                run.font.color.rgb = RGBColor(255, 0, 0)
        # ページ区切り（最後のページは入れなくてもよい場合は条件を調整）
        if page_num < len(ocr_result.pages):
            doc.add_page_break()

    # Wordファイルを保存
    word_filename = f"{base_filename}_ocr_result.docx"
    doc.save(word_filename)

    return all_pages_data


def create_pdf_from_ocr(input_pdf, output_path):
    # 入力PDFのファイル名（拡張子なし）を取得
    base_filename = os.path.splitext(os.path.basename(input_pdf))[0]
    output_dir = os.path.dirname(output_path)

    # Azure OCRで解析
    ocr_result = analyze_pdf(input_pdf)

    # OCR結果を保存
    result_base_path = os.path.join(output_dir, base_filename)
    save_ocr_results(ocr_result, result_base_path)

    # PDFを画像に変換（全ページ）
    images = convert_from_path(input_pdf)

    # 複数ページのPDFを作成
    c = canvas.Canvas(output_path)

    print("PDFを生成中...")
    # 各ページを処理
    for page_num, (image, page) in tqdm(
        enumerate(zip(images, ocr_result.pages)),
        total=len(ocr_result.pages),
        desc="PDF生成",
    ):
        # 一時的に画像を保存
        temp_image_path = f"temp_image_{page_num}.jpg"
        image.save(temp_image_path, "JPEG")

        # 画像のサイズを取得
        width, height = image.size

        # ページサイズを設定
        c.setPageSize((width, height))

        # 変換した画像を背景として描画
        c.drawImage(temp_image_path, 0, 0, width=width, height=height)

        # OCRデータから各テキストの位置を取得して描画
        for word in page.words:
            if word.polygon:
                # ポリゴン座標の取得と変換（スケール調整）
                polygon = word.polygon
                scale = 203  # スケール係数

                # 座標変換（PDFの座標系に合わせる）
                x1, y1 = float(polygon[0].x) * scale, float(polygon[0].y) * scale
                x2, y2 = float(polygon[1].x) * scale, float(polygon[1].y) * scale
                _x3, y3 = float(polygon[2].x) * scale, float(polygon[2].y) * scale
                _x4, y4 = float(polygon[3].x) * scale, float(polygon[3].y) * scale

                # PDFの座標系（下が原点）に変換
                y1 = height - y1
                y2 = height - y2
                y3 = height - y3
                y4 = height - y4

                # テキストの描画位置と文字サイズを計算
                text_height = abs(float(polygon[2].y) - float(polygon[0].y)) * scale

                # confidenceに基づいて色を設定
                if float(word.confidence) < 0.95:
                    c.setFillColor(red)
                else:
                    c.setFillColor(blue)

                # フォントサイズを設定して描画（高さの70%に設定）
                font_size = text_height * 0.7
                c.setFont("BIZUDPMincho", font_size)

                # テキストを矩形の中に配置
                text_width = c.stringWidth(word.content, "BIZUDPMincho", font_size)
                box_width = x2 - x1
                x_offset = (box_width - text_width) / 2
                y_offset = text_height * 0.2

                # テキストを描画
                c.drawString(x1 + x_offset, y1 + y_offset, word.content)

        # 一時ファイルの削除
        os.remove(temp_image_path)

        # 次のページを追加（最後のページ以外）
        if page_num < len(ocr_result.pages) - 1:
            c.showPage()

    # PDFを保存
    c.save()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='PDFファイルからOCRを実行し、テキストを抽出します')
    parser.add_argument('input_pdf', help='入力PDFファイルのパス')
    parser.add_argument('--output_dir', default='output', help='出力ディレクトリのパス（デフォルト: output）')
    
    args = parser.parse_args()

    # 出力ディレクトリが存在しない場合は作成
    os.makedirs(args.output_dir, exist_ok=True)

    # 入力ファイル名から拡張子を除いたベース名を取得
    base_name = os.path.splitext(os.path.basename(args.input_pdf))[0]

    # 出力ファイル名を設定
    output_filename = os.path.join(args.output_dir, f"{base_name}_output.pdf")
    create_pdf_from_ocr(args.input_pdf, output_filename)
